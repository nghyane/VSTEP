from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
CAPSTONE = ROOT / "docs" / "capstone"
TEMPLATE = CAPSTONE / "Template" / "Report6_Software User Guides.docx"
SOURCE = CAPSTONE / "reports" / "report6-software-user-guides.md"
DEFAULT_OUTPUT = CAPSTONE / "reports" / "report6-software-user-guides.docx"
FALLBACK_TEMPLATE = CAPSTONE / "reports" / "report6-software-user-guides-final-v5.docx"
PROJECT_TITLE = "An Adaptive VSTEP Preparation System with Comprehensive Skill Assessment and Personalized Learning Support"
COVER_DATE = "- Ho Chi Minh, May 2026 -"

IMAGE_RE = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<path>.+)\)")
TABLE_HEADER_FILL = "FFE8E1"
TOC_PAGE_NUMBERS = {
    "I. Record of Changes": 3,
    "II. Release Package & User Guides": 3,
    "1. Deliverable Package": 3,
    "2. Installation Guides": 4,
    "2.1 System Requirements": 4,
    "2.1.1 Hardware Requirements": 4,
    "2.1.2 Software Requirements": 4,
    "2.2 Installation Instruction": 5,
    "3. User Manual": 9,
    "3.1 Overview": 9,
    "3.2 Workflow 1": 10,
    "3.3 Workflow 2": 26,
    "3.4 Workflow 3": 53,
    "3.5 Workflow 4": 60,
    "3.6 Workflow 5": 70,
}


def keep_template_cover(document: Document) -> None:
    body = document._body._element
    start = None
    for paragraph in document.paragraphs:
        if "Hanoi" in paragraph.text or "Ho Chi Minh" in paragraph.text:
            start = paragraph._element
    if start is None:
        return

    removing = False
    for child in list(body):
        if removing and child.tag != qn("w:sectPr"):
            body.remove(child)
        if child is start:
            removing = True


def replace_cover_project_title(document: Document) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().lower() == "capstone project report":
            replaced = False
            for run in paragraph.runs:
                if run.text.strip():
                    run.text = PROJECT_TITLE if not replaced else ""
                    replaced = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return


def replace_cover_date(document: Document) -> None:
    for paragraph in document.paragraphs:
        if "2019" in paragraph.text or "Hanoi" in paragraph.text:
            replaced = False
            for run in paragraph.runs:
                if run.text.strip():
                    run.text = COVER_DATE if not replaced else ""
                    replaced = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return


def extract_table_of_contents(source_lines: list[str]) -> list[str]:
    try:
        start = source_lines.index("## Table of Contents") + 1
        end = source_lines.index("## I. Record of Changes")
    except ValueError:
        return []

    entries: list[str] = []
    for line in source_lines[start:end]:
        entry = line.strip()
        if not entry:
            continue
        workflow = re.match(r"^(3\.\d+ Workflow \d+):", entry)
        entries.append(workflow.group(1) if workflow else entry)
    return entries


def add_table_of_contents(document: Document, entries: list[str]) -> None:
    title = document.add_paragraph()
    run = title.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    for entry in entries:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        if re.match(r"^[IVX]+\.\s+", entry):
            paragraph.paragraph_format.left_indent = Inches(0)
        elif re.match(r"^\d+\.\d+", entry):
            paragraph.paragraph_format.left_indent = Inches(0.45)
        elif re.match(r"^\d+\.\s+", entry):
            paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.add_run(f"{entry}\t{TOC_PAGE_NUMBERS.get(entry, '')}")
    document.add_page_break()


def add_run_text(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(9)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph()
    add_run_text(paragraph, text)
    return paragraph


def add_heading(document: Document, text: str, level: int) -> None:
    if text == "Purpose" or text.startswith("Step "):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        return

    if re.match(r"^2\.2\.\d+", text):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        return

    workflow = re.match(r"^(3\.\d+ Workflow \d+):", text)
    if workflow:
        document.add_heading(workflow.group(1), level=3)
        return

    if re.match(r"^[IVX]+\.\s+", text):
        document.add_heading(text, level=1)
        return
    if re.match(r"^\d+\.\d+", text):
        document.add_heading(text, level=3)
        return
    if re.match(r"^\d+\.\s+", text):
        document.add_heading(text, level=2)
        return
    if level >= 3:
        document.add_heading(text, level=3)
        return
    document.add_heading(text, level=1)


def add_code_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9)


def add_image(document: Document, markdown_dir: Path, alt: str, image_path: str) -> list[str]:
    resolved = (markdown_dir / image_path).resolve()
    if not resolved.exists():
        return [str(resolved)]

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    image_width = Inches(2.6) if "anh_mobile" in image_path else Inches(6.2)
    run.add_picture(str(resolved), width=image_width)

    if alt:
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run(alt)
        caption_run.italic = True
        caption_run.font.size = Pt(9)
    return []


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Normal Table"
    set_table_borders(table)

    for row_index, row in enumerate(rows):
        cells = table.rows[row_index].cells
        for column_index, value in enumerate(row):
            if column_index >= len(cells):
                continue
            cells[column_index].text = value.strip()
            if row_index == 0:
                set_cell_shading(cells[column_index], TABLE_HEADER_FILL)
            for paragraph in cells[column_index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    if row_index == 0:
                        run.bold = True


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip().strip("|")
        cells = [cell.strip() for cell in raw.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def build_docx(output: Path = DEFAULT_OUTPUT) -> None:
    template = TEMPLATE if TEMPLATE.exists() else FALLBACK_TEMPLATE
    document = Document(str(template))
    replace_cover_project_title(document)
    replace_cover_date(document)
    keep_template_cover(document)

    markdown_dir = SOURCE.parent
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    add_table_of_contents(document, extract_table_of_contents(source_lines))
    try:
        start_index = source_lines.index("## I. Record of Changes")
    except ValueError:
        start_index = 0
    lines = source_lines[start_index:]
    missing_images: list[str] = []
    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        image = IMAGE_RE.fullmatch(stripped)
        if image:
            missing_images.extend(add_image(document, markdown_dir, image.group("alt"), image.group("path")))
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            add_heading(document, heading.group(2), len(heading.group(1)))
            index += 1
            continue

        if stripped.startswith("- "):
            text = stripped[2:]
            if text.startswith("Ho Chi Minh"):
                paragraph = add_paragraph(document, text)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_paragraph(document, f"- {text}")
            index += 1
            continue

        add_paragraph(document, stripped)
        index += 1

    if missing_images:
        raise FileNotFoundError("Missing images:\n" + "\n".join(missing_images))

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    print(f"Generated {output}")


if __name__ == "__main__":
    build_docx(Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else DEFAULT_OUTPUT)
