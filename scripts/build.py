#!/usr/bin/env python3
"""
MD to DOCX Builder - Hybrid approach with config-driven styles.
"""

from typing import cast, Any, Dict, TYPE_CHECKING
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree  # type: ignore[import-untyped]
from pathlib import Path
import re
import yaml

if TYPE_CHECKING:
    from docx.document import Document as DocType


# XML Namespaces for OOXML
NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}


class EndnoteManager:
    """Manages native Word endnotes using OOXML manipulation.
    
    Word endnotes require:
    1. An endnotes.xml part in the document package
    2. w:endnoteReference elements in the document body
    3. w:endnote elements in endnotes.xml with matching IDs
    
    Endnotes appear at END OF DOCUMENT (not bottom of each page like footnotes).
    """
    
    def __init__(self, doc: Any, endnote_definitions: Dict[str, str]):
        """Initialize endnote manager.
        
        Args:
            doc: The python-docx Document object
            endnote_definitions: Dict mapping endnote ID (str) to endnote text
        """
        self.doc = doc
        self.definitions = endnote_definitions
        self._endnotes_part = None
        self._endnotes_element = None
        self._next_id = 1  # ID 0 is reserved for separator, -1 for continuation
        self._id_mapping: Dict[str, int] = {}  # Map markdown IDs to Word IDs
        
        if endnote_definitions:
            self._initialize_endnotes_part()
    
    def _initialize_endnotes_part(self):
        """Create or get the endnotes.xml part in the document."""
        document_part = self.doc.part
        
        # Check if endnotes part already exists
        for rel in document_part.rels.values():
            if rel.reltype == RT.ENDNOTES:
                self._endnotes_part = rel.target_part
                self._endnotes_element = self._endnotes_part._element
                # Find highest existing ID
                for endnote in self._endnotes_element.findall('.//w:endnote', NSMAP):
                    eid = int(endnote.get(qn('w:id'), '0'))
                    if eid >= self._next_id:
                        self._next_id = eid + 1
                return
        
        # Create new endnotes part
        # Build the endnotes.xml content with separator endnotes
        endnotes_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
    <w:endnote w:type="separator" w:id="-1">
        <w:p>
            <w:r>
                <w:separator/>
            </w:r>
        </w:p>
    </w:endnote>
    <w:endnote w:type="continuationSeparator" w:id="0">
        <w:p>
            <w:r>
                <w:continuationSeparator/>
            </w:r>
        </w:p>
    </w:endnote>
</w:endnotes>'''
        
        # Parse and create the endnotes part
        self._endnotes_element = etree.fromstring(endnotes_xml.encode('utf-8'))
        
        # Create a new part for endnotes
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        
        # Create endnotes part
        endnotes_part_uri = PackURI('/word/endnotes.xml')
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml'
        
        # Create a custom part class to hold our endnotes
        class EndnotesPart(Part):
            def __init__(self, partname, content_type, element, package):
                super().__init__(partname, content_type, package=package)
                self._element = element
            
            @property
            def blob(self):
                return etree.tostring(self._element, xml_declaration=True, 
                                     encoding='UTF-8', standalone=True)
        
        self._endnotes_part = EndnotesPart(
            endnotes_part_uri, 
            content_type, 
            self._endnotes_element,
            document_part.package
        )
        
        # Add relationship from document to endnotes
        document_part.relate_to(self._endnotes_part, RT.ENDNOTES)
    
    def add_endnote_reference(self, paragraph, markdown_id: str) -> bool:
        """Add a clickable endnote reference to a paragraph.
        
        Shows [N] format inline in black (IEEE standard).
        The w:endnoteReference provides clickable jump-to-endnote functionality.
        
        Args:
            paragraph: The python-docx Paragraph object
            markdown_id: The endnote ID from markdown (e.g., "5" from [^5])
            
        Returns:
            True if endnote was added, False if definition not found
        """
        if markdown_id not in self.definitions:
            # Fallback if no definition - use IEEE standard styling
            run = paragraph.add_run(f"[{markdown_id}]")
            run.font.superscript = False  # Inline, not superscript
            run.font.size = Pt(11)  # Same as body text
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            run.font.name = 'Calibri'
            return False
        
        # Get or assign Word endnote ID
        if markdown_id not in self._id_mapping:
            word_id = self._next_id
            self._next_id += 1
            self._id_mapping[markdown_id] = word_id
            # Add endnote content to endnotes.xml
            self._add_endnote_content(word_id, self.definitions[markdown_id])
        else:
            word_id = self._id_mapping[markdown_id]
        
        # Add visible [N] text inline (IEEE standard)
        bracket_run = paragraph.add_run(f'[{word_id}]')
        bracket_run.font.superscript = False  # Inline, not superscript
        bracket_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        bracket_run.font.size = Pt(11)  # Same as body text
        bracket_run.font.name = 'Calibri'
        
        # Create hidden endnote reference for clickable functionality
        # The w:endnoteReference element provides the link to endnote
        run = paragraph.add_run()
        run_element = run._r
        
        # Add run properties to hide the auto-generated number
        rPr = run_element.get_or_add_rPr()
        # Set vanish property to hide the default endnote number
        vanish = OxmlElement('w:vanish')
        rPr.append(vanish)
        
        # Add the endnote reference element (hidden, but provides jump functionality)
        endnote_ref = OxmlElement('w:endnoteReference')
        endnote_ref.set(qn('w:id'), str(word_id))
        run_element.append(endnote_ref)
        
        return True
    
    def _add_endnote_content(self, word_id: int, text: str):
        """Add endnote content to endnotes.xml."""
        if self._endnotes_element is None:
            return
        
        # Create the endnote element
        endnote = etree.SubElement(self._endnotes_element, qn('w:endnote'))
        endnote.set(qn('w:id'), str(word_id))
        
        # Create paragraph
        p = etree.SubElement(endnote, qn('w:p'))
        
        # Add paragraph properties with endnote style
        pPr = etree.SubElement(p, qn('w:pPr'))
        pStyle = etree.SubElement(pPr, qn('w:pStyle'))
        pStyle.set(qn('w:val'), 'EndnoteText')
        
        # First run: endnote reference mark (the number that appears in endnote area)
        r1 = etree.SubElement(p, qn('w:r'))
        rPr1 = etree.SubElement(r1, qn('w:rPr'))
        rStyle1 = etree.SubElement(rPr1, qn('w:rStyle'))
        rStyle1.set(qn('w:val'), 'EndnoteReference')
        endnote_ref = etree.SubElement(r1, qn('w:endnoteRef'))
        
        # Second run: space after number
        r2 = etree.SubElement(p, qn('w:r'))
        t2 = etree.SubElement(r2, qn('w:t'))
        t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t2.text = ' '
        
        # Third run: endnote text content
        r3 = etree.SubElement(p, qn('w:r'))
        t3 = etree.SubElement(r3, qn('w:t'))
        t3.text = text
    
    def finalize(self):
        """Ensure endnotes are properly saved. Called after document is built."""
        # The endnotes part is already connected via relationship
        # Just ensure the element is up to date
        pass


def parse_footnote_definitions(md_text: str) -> Dict[str, str]:
    """Extract footnote definitions from markdown text.
    
    Looks for patterns like:
    [^1]: This is the footnote text.
    [^2]: Another footnote that can span
        multiple lines with indentation.
    
    Returns:
        Dict mapping footnote ID (str) to footnote text
    """
    definitions = {}
    lines = md_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        # Match footnote definition start: [^N]: text
        match = re.match(r'^\[\^(\d+)\]:\s*(.*)$', line)
        if match:
            footnote_id = match.group(1)
            text_parts = [match.group(2)]
            
            # Check for continuation lines (indented)
            i += 1
            while i < len(lines):
                next_line = lines[i]
                # Continuation lines are indented (4 spaces or tab)
                if next_line.startswith('    ') or next_line.startswith('\t'):
                    text_parts.append(next_line.strip())
                    i += 1
                elif next_line.strip() == '':
                    # Empty line might be part of multi-paragraph footnote
                    i += 1
                else:
                    break
            
            definitions[footnote_id] = ' '.join(text_parts).strip()
        else:
            i += 1
    
    return definitions


def load_style_config():
    """Load typography configuration from YAML."""
    config_path = Path(__file__).parent / "styles.yaml"
    with open(config_path, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f)


def ensure_styles(doc, config):
    """Ensure all styles from config exist in document.
    
    Creates missing styles, updates existing ones from config.
    """
    style_config = config.get('styles', {})
    base_config = config.get('base', {})
    
    for style_name, props in style_config.items():
        # Check if style exists
        try:
            style = doc.styles[style_name]
        except KeyError:
            # Create new style
            if style_name.startswith('Heading'):
                style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                style.base_style = doc.styles['Normal']
            else:
                style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        
        # Apply properties from config
        font = style.font
        pf = style.paragraph_format
        
        # Font properties
        font.name = props.get('font', base_config.get('font', 'Calibri'))
        font.size = Pt(props.get('size', base_config.get('size', 11)))
        
        if props.get('bold'):
            font.bold = True
        if props.get('italic'):
            font.italic = True
        if props.get('color'):
            color_hex = props['color']
            font.color.rgb = RGBColor(
                int(color_hex[0:2], 16),
                int(color_hex[2:4], 16),
                int(color_hex[4:6], 16)
            )
        
        # Paragraph properties
        if 'space_before' in props:
            pf.space_before = Pt(props['space_before'])
        if 'space_after' in props:
            pf.space_after = Pt(props['space_after'])
        if 'line_spacing' in props:
            pf.line_spacing = props['line_spacing']  # Multiple (1.15, 1.5, etc.)
        if props.get('keep_with_next'):
            pf.keep_with_next = True
        if 'left_indent' in props:
            pf.left_indent = Pt(props['left_indent'])
        if 'first_line_indent' in props:
            pf.first_line_indent = Pt(props['first_line_indent'])
        
        # Outline level for headings (used by TOC)
        if 'outline_level' in props:
            # Set outline level via XML
            pPr = style.element.get_or_add_pPr()
            outline = OxmlElement('w:outlineLvl')
            outline.set(qn('w:val'), str(props['outline_level']))
            # Remove existing if any
            existing = pPr.find(qn('w:outlineLvl'))
            if existing is not None:
                pPr.remove(existing)
            pPr.append(outline)


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    """Set cell margins/padding in twips (dxa). 20 twips = 1pt."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(margin_value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_width(table, width_pct=100):
    """Set table width as percentage of page width."""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_pct * 50))  # 5000 = 100%
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

# Style mapping - template style names
STYLE_MAP = {
    1: "Heading 1",
    2: "Heading 2", 
    3: "Heading 3",
    4: "Heading 4",
    5: "Heading 5",
    6: "Heading 6",
    "normal": "Normal",
    "bold_label": "Normal",  # Bold subheadings like "**Quy mô:**"
    "bullet": "List Paragraph",
    "number": "List Paragraph",
    "table": "Table Grid",
    "table_head": "Table Head",
    "table_body": "Normal",
}

# Fallbacks if style not in template
STYLE_FALLBACK = {
    "List Paragraph": "Normal",
    "Normal Table": None,
    "Table Head": "Normal",
    "Table Text small": "Normal",
}

def get_style(doc, style_key):
    """Get style name, falling back if not available."""
    style_name = STYLE_MAP.get(style_key, "Normal")
    style_names = {s.name for s in doc.styles}
    if style_name in style_names:
        return style_name
    return STYLE_FALLBACK.get(style_name, "Normal")

def parse_md(md_text):
    """Parse markdown into structured blocks.
    
    Skips content before the first heading (cover page metadata).
    """
    blocks = []
    lines = md_text.split('\n')
    i = 0
    found_first_heading = False
    
    while i < len(lines):
        line = lines[i]
        
        # Skip content before first heading (cover page in MD)
        if not found_first_heading:
            if line.startswith('# '):
                found_first_heading = True
            else:
                i += 1
                continue
        
        # Headings H1-H6
        if line.startswith('# '):
            blocks.append(('heading', 1, line[2:].strip()))
        elif line.startswith('## '):
            blocks.append(('heading', 2, line[3:].strip()))
        elif line.startswith('### '):
            blocks.append(('heading', 3, line[4:].strip()))
        elif line.startswith('#### '):
            blocks.append(('heading', 4, line[5:].strip()))
        elif line.startswith('##### '):
            blocks.append(('heading', 5, line[6:].strip()))
        elif line.startswith('###### '):
            blocks.append(('heading', 6, line[7:].strip()))
        
        # Bullet list
        elif line.startswith('- ') or line.startswith('* '):
            blocks.append(('bullet', line[2:].strip()))
        
        # Numbered list
        elif re.match(r'^\d+\. ', line):
            blocks.append(('number', re.sub(r'^\d+\. ', '', line).strip()))
        
        # Table
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                if not lines[i].startswith('|--') and not lines[i].startswith('| --'):
                    table_lines.append(lines[i])
                i += 1
            i -= 1
            blocks.append(('table', parse_table(table_lines)))
        
        # Normal paragraph (check for bold label subheading)
        elif line.strip():
            text = line.strip()
            # Detect bold label subheadings like "**Quy mô:**" or "**Tính năng:**"
            if re.match(r'^\*\*[^*]+:\*\*', text):
                blocks.append(('bold_label', text))
            else:
                blocks.append(('paragraph', text))
        
        i += 1
    
    return blocks

def parse_table(lines):
    """Parse markdown table lines into 2D array."""
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    return rows

def add_endnote(paragraph, endnote_id: str, endnote_manager: Any):
    """Add a native Word endnote reference to a paragraph.
    
    Uses EndnoteManager to create clickable endnotes via OOXML.
    Falls back to superscript if manager is None or a dict (legacy).
    """
    if endnote_manager is None:
        # Fallback: just superscript
        run = paragraph.add_run(f"[{endnote_id}]")
        run.font.superscript = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
    elif isinstance(endnote_manager, EndnoteManager):
        # Use native Word endnotes
        endnote_manager.add_endnote_reference(paragraph, endnote_id)
    else:
        # Legacy dict fallback
        run = paragraph.add_run(f"[{endnote_id}]")
        run.font.superscript = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'


def render_inline(paragraph, text, endnote_manager=None):
    """Render inline formatting (bold, italic, endnotes) within a paragraph.
    
    Args:
        paragraph: The python-docx Paragraph object
        text: The text to render with markdown formatting
        endnote_manager: EndnoteManager instance for native endnotes
    """
    # First, handle endnote references [^N]
    # Split by endnote markers
    endnote_pattern = r'(\[\^\d+\])'
    parts = re.split(endnote_pattern, text)
    
    for part in parts:
        # Check if this is an endnote reference
        endnote_match = re.match(r'\[\^(\d+)\]', part)
        if endnote_match:
            endnote_id = endnote_match.group(1)
            add_endnote(paragraph, endnote_id, endnote_manager)
        elif part:
            # Process bold/italic for non-endnote parts
            render_inline_formatting(paragraph, part)


def render_inline_formatting(paragraph, text):
    """Render bold and italic formatting within text."""
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)  # Same size as normal, bold weight is enough
            run.font.name = 'Calibri'
        elif part:
            # Handle italic within non-bold parts
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*') and not ip.startswith('**'):
                    run = paragraph.add_run(ip[1:-1])
                    run.italic = True
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
                elif ip:
                    run = paragraph.add_run(ip)
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'

def find_content_start_index(doc):
    """Find where content should start (after cover/TOC).
    
    Returns the index AFTER TOC section ends.
    Strategy:
    1. Look for existing content markers ("I. Record of Changes", etc.)
    2. If not found, look for end of TOC section (after "Table of Contents" heading)
    3. If TOC found, return index after TOC instructions
    """
    # First, look for existing content section markers
    content_markers = [
        "I. Record of Changes",
        "Record of Changes", 
        "I.",  # Any section starting with Roman numeral I.
    ]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for marker in content_markers:
            if text.startswith(marker):
                return i
    
    # If no content marker, find end of TOC section
    # Look for "Table of Contents" heading, then find last TOC-related paragraph
    toc_index = -1
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "Table of Contents" in text:
            toc_index = i
            break
    
    if toc_index >= 0:
        # Find the last paragraph that's part of TOC (instructions, etc.)
        # Look for paragraphs after TOC that contain TOC-related text
        last_toc_para = toc_index
        for i in range(toc_index + 1, len(doc.paragraphs)):
            text = doc.paragraphs[i].text.strip()
            if text and ("TOC" in text or "Table of Contents" in text or 
                        "Insert Table" in text or "Update Field" in text):
                last_toc_para = i
            elif text and not text.startswith("("):
                # Found a non-empty, non-instruction paragraph
                break
        
        # Return index after last TOC paragraph
        return last_toc_para + 1
    
    # If nothing found, return -1 (append at end)
    return -1

def clear_content_after_toc(doc, start_index):
    """Remove all body content from start_index onwards.
    
    This clears paragraphs and tables after the TOC section,
    preserving cover page, TOC, and section properties (page layout).
    """
    if start_index < 0:
        return  # No marker found, don't clear anything
    
    # Get the body element (parent of paragraphs)
    body = doc.element.body
    
    if start_index >= len(doc.paragraphs):
        return
    
    # Get the XML element of the paragraph at start_index
    start_para_element = doc.paragraphs[start_index]._element
    
    # Find all elements after this paragraph and remove them
    # BUT preserve sectPr (section properties) which contains page layout
    from docx.oxml.ns import qn
    SECT_PR_TAG = qn('w:sectPr')
    
    found_start = False
    elements_to_remove = []
    
    for element in body:
        if element == start_para_element:
            found_start = True
            # Include this element too (it will be re-added from MD)
            elements_to_remove.append(element)
        elif found_start:
            # Don't remove section properties
            if element.tag != SECT_PR_TAG:
                elements_to_remove.append(element)
    
    for element in elements_to_remove:
        body.remove(element)

def build_docx(md_path, template_path, output_path):
    """Build DOCX from MD using template styles.
    
    Preserves cover page and TOC from template, only adds/replaces content after.
    """
    
    # Load template
    doc = Document(str(template_path))
    
    # Load style config and ensure all styles exist
    style_config = load_style_config()
    ensure_styles(doc, style_config)
    print("✅ Styles ensured from config")
    
    # Parse markdown
    md_text = Path(md_path).read_text(encoding='utf-8')
    blocks = parse_md(md_text)
    
    # Parse endnote definitions and create manager
    endnote_defs = parse_footnote_definitions(md_text)
    endnote_manager = EndnoteManager(doc, endnote_defs) if endnote_defs else None
    if endnote_defs:
        print(f"📝 Found {len(endnote_defs)} endnote definition(s)")
    
    # Find where content should start (after cover/TOC)
    content_start = find_content_start_index(doc)
    
    if content_start >= 0 and content_start < len(doc.paragraphs):
        para_text = doc.paragraphs[content_start].text.strip()
        if para_text:
            print(f"📍 Found content at paragraph {content_start}: '{para_text[:50]}...'")
        else:
            print(f"📍 Content area starts at paragraph {content_start} (after TOC)")
        # Clear existing content after TOC (if any sample content exists)
        clear_content_after_toc(doc, content_start)
    elif content_start >= len(doc.paragraphs):
        print(f"📍 TOC ends at paragraph {content_start - 1}, appending content after")
    else:
        print("📍 No TOC marker found, appending at end of template")
    
    # Now add content from markdown
    for block in blocks:
        block_type = block[0]
        
        if block_type == 'heading':
            level = block[1]
            text = block[2]
            style_name = get_style(doc, level)
            p = doc.add_paragraph(style=style_name)
            render_inline(p, text, endnote_manager)
            
        elif block_type == 'paragraph':
            text = block[1]
            p = doc.add_paragraph(style=get_style(doc, "normal"))
            render_inline(p, text, endnote_manager)
        
        elif block_type == 'bold_label':
            # Bold subheading like "**Quy mô:**" - add extra spacing before
            text = block[1]
            p = doc.add_paragraph(style=get_style(doc, "bold_label"))
            p.paragraph_format.space_before = Pt(6)  # Extra spacing before
            render_inline(p, text, endnote_manager)
            
        elif block_type == 'bullet':
            text = block[1]
            style_name = get_style(doc, "bullet")
            p = doc.add_paragraph(style=style_name)
            # Always add bullet with proper formatting
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-18)  # Hanging indent
            p.add_run("• ")
            render_inline(p, text, endnote_manager)
            
        elif block_type == 'number':
            text = block[1]
            p = doc.add_paragraph(style=get_style(doc, "number"))
            # Match bullet indent for consistency
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-18)  # Hanging indent
            render_inline(p, text, endnote_manager)
            
        elif block_type == 'table':
            rows = block[1]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table_style = get_style(doc, "table")
                if table_style:
                    try:
                        table.style = table_style
                    except:
                        pass  # Style may not exist
                
                # Set table width to 100% for Google Docs/365 compatibility
                set_table_width(table, 100)
                
                # Get cell styles
                head_style = get_style(doc, "table_head")
                body_style = get_style(doc, "table_body")
                
                for i, row_data in enumerate(rows):
                    table_row = table.rows[i]
                    # Set row height rule to auto (minimum)
                    table_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                    table_row.height = Pt(20)  # Minimum height
                    
                    for j, cell_text in enumerate(row_data):
                        cell = table.cell(i, j)
                        cell.text = ""
                        set_cell_margins(cell)  # Add padding (2pt top/bottom, 4pt left/right)
                        p = cell.paragraphs[0]
                        # Apply "Table Head" to first row, "Normal" to others
                        cell_style = head_style if i == 0 else body_style
                        if cell_style:
                            try:
                                p.style = cell_style
                            except:
                                pass
                        # Apply background color to header row
                        if i == 0:
                            set_cell_shading(cell, "FFE8E1")  # Light pink header
                        # Render text
                        render_inline(p, cell_text, endnote_manager)
                        # Apply font from style (runs don't inherit automatically)
                        style_obj = doc.styles[cell_style] if cell_style else None
                        if style_obj:
                            # Cast to Any to avoid LSP errors - python-docx types are incomplete
                            style_font = cast(Any, style_obj).font if hasattr(style_obj, 'font') else None
                            for run in p.runs:
                                if style_font and style_font.name:
                                    run.font.name = style_font.name
                                if style_font and style_font.size:
                                    run.font.size = style_font.size
                                if i == 0:
                                    run.bold = True
    
    # Finalize footnotes
    if endnote_manager:
        endnote_manager.finalize()
    
    # Save
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✅ Created: {output_path}")

def main():
    import sys
    import argparse
    
    project_root = Path(__file__).parent.parent
    template = project_root / "docs/capstone/templates/fpt-report1-template.docx"
    output_path = project_root / "docs/capstone/output"
    
    # Support both VI and ENG folders
    vi_path = project_root / "docs/capstone/reports/VI"
    eng_path = project_root / "docs/capstone/reports/ENG"
    
    # Parse arguments
    parser = argparse.ArgumentParser(description='Build DOCX from MD')
    parser.add_argument('name', nargs='?', default='report1-project-introduction',
                        help='Markdown file name (without extension)')
    parser.add_argument('-l', '--lang', choices=['VI', 'ENG', 'all'], default='VI',
                        help='Language folder: VI, ENG, or all (default: VI)')
    parser.add_argument('-o', '--output', default=str(output_path),
                        help=f'Output directory (default: {output_path})')
    
    args = parser.parse_args()
    
    def build_single(md_folder, lang):
        """Build DOCX for a specific language folder."""
        md_file = md_folder / f"{args.name}.md"
        if md_file.exists():
            output_file = Path(args.output) / f"{args.name}.docx"
            build_docx(md_file, template, output_file)
            return True
        return False
    
    # Build according to language flag
    built = []
    if args.lang in ['VI', 'all']:
        if build_single(vi_path, 'VI'):
            built.append('VI')
    if args.lang in ['ENG', 'all']:
        if build_single(eng_path, 'ENG'):
            built.append('ENG')
    
    if not built:
        print(f"❌ File not found: {args.name}.md in VI or ENG folders")
        sys.exit(1)
    
    print(f"✅ Built: {', '.join(built)} versions")

if __name__ == "__main__":
    main()
