#!/usr/bin/env python3
"""Create minimal FPT template - keep cover, headers, styles, remove sample content."""

from docx import Document
from pathlib import Path

def create_minimal_template():
    src = Path("docs/capstone/templates/fpt-report1-template.docx")
    dst = Path("docs/capstone/templates/fpt-minimal-template.docx")
    
    doc = Document(str(src))
    
    # Find "II. Project Introduction" and remove everything after it except structure
    found_section_ii = False
    paragraphs_to_clear = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Start clearing after "II. Project Introduction"
        if "II." in text and "Project Introduction" in text:
            found_section_ii = True
            continue
        
        if found_section_ii:
            # Keep headings (usually have style Heading 1/2/3 or start with numbers)
            style_name = para.style.name if para.style else ""
            is_heading = "Heading" in style_name
            
            # Check if it's a numbered heading like "1. Overview" or "3.1 System"
            is_numbered = any(text.startswith(f"{n}.") or text.startswith(f"{n} ") for n in range(1, 10))
            is_sub_numbered = any(text.startswith(f"{n}.{m}") for n in range(1, 10) for m in range(1, 10))
            
            if is_heading or is_numbered or is_sub_numbered:
                # Keep heading but clear any trailing sample text
                continue
            
            # Mark for clearing if contains sample markers or is regular paragraph
            if "<<" in text or "Sample" in text or (text and not is_heading):
                paragraphs_to_clear.append(para)
    
    # Clear content (keep paragraph for structure, just empty the text)
    for para in paragraphs_to_clear:
        for run in para.runs:
            run.text = ""
    
    # Clear table data rows (keep headers)
    for table in doc.tables:
        # Keep first row (header), remove others
        while len(table.rows) > 1:
            tr = table.rows[-1]._tr
            table._tbl.remove(tr)
    
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    print(f"✅ Created: {dst}")

if __name__ == "__main__":
    create_minimal_template()
